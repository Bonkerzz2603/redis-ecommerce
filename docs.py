"""
Script tạo tài liệu DOCX cho project Retailrocket Recommender.

Chạy script này để sinh file tài liệu Word (.docx) mô tả chi tiết
về project, cấu trúc, cách sử dụng và hướng dẫn.

Cách sử dụng:
    python generate_docs.py
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os


def create_documentation():
    """Tạo file tài liệu DOCX cho project."""
    doc = Document()

    # ======= Thiết lập style =======
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)

    # ======= TRANG BÌA =======
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('BÁO CÁO ĐỒ ÁN')
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('RETAILROCKET RECOMMENDER SYSTEM')
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0, 102, 153)

    doc.add_paragraph()

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run('Hệ thống phân tích và gợi ý sản phẩm\n'
                       'cho nền tảng thương mại điện tử')
    run.font.size = Pt(14)
    run.italic = True

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('Công nghệ sử dụng: Python, Pandas, Matplotlib, Tkinter\n'
                       'Nguồn dữ liệu: Kaggle - Retailrocket E-Commerce Dataset')
    run.font.size = Pt(12)

    doc.add_page_break()

    # ======= MỤC LỤC =======
    doc.add_heading('MỤC LỤC', level=1)
    toc_items = [
        '1. Giới thiệu',
        '2. Mô tả dữ liệu',
        '3. Cấu trúc project',
        '4. Yêu cầu hệ thống',
        '5. Hướng dẫn cài đặt',
        '6. Hướng dẫn sử dụng',
        '7. Các chức năng chính',
        '8. Mô tả kỹ thuật',
        '9. Kết luận',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ======= 1. GIỚI THIỆU =======
    doc.add_heading('1. Giới thiệu', level=1)
    doc.add_paragraph(
        'Retailrocket Recommender System là một ứng dụng phân tích dữ liệu '
        'thương mại điện tử (E-Commerce) được xây dựng dựa trên dataset '
        'Retailrocket từ Kaggle. Ứng dụng cung cấp các công cụ để tải, '
        'phân tích và trực quan hóa dữ liệu hành vi người dùng trên '
        'nền tảng thương mại điện tử.'
    )
    doc.add_paragraph(
        'Dự án được phát triển từ một Jupyter Notebook ban đầu và được '
        'tổ chức lại thành một project Python có cấu trúc rõ ràng, '
        'bao gồm giao diện đồ họa (GUI) sử dụng thư viện Tkinter, '
        'giúp người dùng có thể tương tác trực quan với dữ liệu mà '
        'không cần kiến thức lập trình.'
    )

    doc.add_heading('1.1. Mục tiêu', level=2)
    objectives = [
        'Tải và tiền xử lý dữ liệu từ các file CSV của Retailrocket.',
        'Phân tích thống kê mô tả cho các tập dữ liệu.',
        'Trực quan hóa dữ liệu bằng các biểu đồ phân bổ, tương quan, scatter/density.',
        'Cung cấp giao diện đồ họa thân thiện cho việc khám phá dữ liệu.',
        'Tổ chức mã nguồn thành cấu trúc project Python chuẩn.',
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')

    # ======= 2. MÔ TẢ DỮ LIỆU =======
    doc.add_heading('2. Mô tả dữ liệu', level=1)
    doc.add_paragraph(
        'Dataset Retailrocket chứa dữ liệu hành vi người dùng trên một '
        'nền tảng thương mại điện tử, bao gồm 4 file CSV:'
    )

    # Bảng mô tả files
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Shading Accent 1'
    headers = ['File', 'Mô tả', 'Nội dung chính']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    data_files = [
        ['events.csv', 'Sự kiện người dùng',
         'timestamp, visitorid, event (view/addtocart/transaction), itemid, transactionid'],
        ['category_tree.csv', 'Cây phân loại sản phẩm',
         'categoryid, parentid'],
        ['item_properties_part1.csv', 'Thuộc tính sản phẩm (phần 1)',
         'timestamp, itemid, property, value'],
        ['item_properties_part2.csv', 'Thuộc tính sản phẩm (phần 2)',
         'timestamp, itemid, property, value'],
    ]
    for row_idx, row_data in enumerate(data_files, 1):
        for col_idx, text in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = text

    doc.add_paragraph()
    doc.add_heading('2.1. Chi tiết file events.csv', level=2)
    doc.add_paragraph(
        'File events.csv là file dữ liệu chính, chứa thông tin về các sự kiện '
        'tương tác của người dùng với sản phẩm. Các loại sự kiện bao gồm:'
    )
    event_types = [
        'view: Người dùng xem sản phẩm (chiếm đa số)',
        'addtocart: Người dùng thêm sản phẩm vào giỏ hàng',
        'transaction: Người dùng thực hiện giao dịch mua hàng',
    ]
    for et in event_types:
        doc.add_paragraph(et, style='List Bullet')

    # ======= 3. CẤU TRÚC PROJECT =======
    doc.add_heading('3. Cấu trúc project', level=1)
    doc.add_paragraph('Project được tổ chức theo cấu trúc sau:')

    structure = [
        'redis-ecommerce-redis-project/',
        '├── main.py                    # File chạy chính (entry point)',
        '├── requirements.txt           # Danh sách dependencies',
        '├── README.md                  # Tài liệu hướng dẫn',
        '├── generate_docs.py           # Script tạo tài liệu DOCX',
        '├── data/                      # Thư mục chứa dữ liệu CSV',
        '│   ├── events.csv',
        '│   ├── category_tree.csv',
        '│   ├── item_properties_part1.csv',
        '│   └── item_properties_part2.csv',
        '├── src/                       # Mã nguồn chính',
        '│   ├── __init__.py',
        '│   ├── data_loader.py         # Module tải dữ liệu',
        '│   ├── analysis.py            # Module phân tích thống kê',
        '│   ├── visualization.py       # Module trực quan hóa',
        '│   └── gui.py                 # Module giao diện đồ họa',
        '├── output/                    # Thư mục lưu kết quả',
        '└── *.ipynb                    # Notebook gốc',
    ]
    for line in structure:
        p = doc.add_paragraph(line)
        p.style.font.name = 'Consolas'
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    doc.add_paragraph()

    # ======= 4. YÊU CẦU HỆ THỐNG =======
    doc.add_heading('4. Yêu cầu hệ thống', level=1)

    doc.add_heading('4.1. Phần mềm', level=2)
    sw_reqs = [
        'Python 3.8 trở lên',
        'pip (Python package manager)',
        'Hệ điều hành: Windows / macOS / Linux',
    ]
    for req in sw_reqs:
        doc.add_paragraph(req, style='List Bullet')

    doc.add_heading('4.2. Thư viện Python', level=2)
    table2 = doc.add_table(rows=7, cols=3)
    table2.style = 'Light Shading Accent 1'
    lib_headers = ['Thư viện', 'Phiên bản', 'Mục đích']
    for i, h in enumerate(lib_headers):
        cell = table2.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    libs = [
        ['pandas', '>= 1.5.0', 'Xử lý và phân tích dữ liệu'],
        ['numpy', '>= 1.23.0', 'Tính toán số học'],
        ['matplotlib', '>= 3.6.0', 'Trực quan hóa dữ liệu'],
        ['seaborn', '>= 0.12.0', 'Biểu đồ thống kê nâng cao'],
        ['scikit-learn', '>= 1.1.0', 'Machine learning và tiền xử lý'],
        ['python-docx', '>= 1.0.0', 'Tạo tài liệu Word'],
    ]
    for row_idx, row_data in enumerate(libs, 1):
        for col_idx, text in enumerate(row_data):
            table2.rows[row_idx].cells[col_idx].text = text

    # ======= 5. HƯỚNG DẪN CÀI ĐẶT =======
    doc.add_heading('5. Hướng dẫn cài đặt', level=1)

    doc.add_heading('Bước 1: Tạo môi trường ảo', level=2)
    doc.add_paragraph('python -m venv venv')
    doc.add_paragraph('Windows: venv\\Scripts\\activate')
    doc.add_paragraph('macOS/Linux: source venv/bin/activate')

    doc.add_heading('Bước 2: Cài đặt dependencies', level=2)
    doc.add_paragraph('pip install -r requirements.txt')

    doc.add_heading('Bước 3: Tải dữ liệu', level=2)
    doc.add_paragraph(
        'Tải dataset Retailrocket từ Kaggle '
        '(https://www.kaggle.com/datasets/retailrocket/ecommerce-dataset) '
        'và giải nén các file CSV vào thư mục data/.'
    )

    doc.add_heading('Bước 4: Chạy ứng dụng', level=2)
    doc.add_paragraph('python main.py')

    # ======= 6. HƯỚNG DẪN SỬ DỤNG =======
    doc.add_heading('6. Hướng dẫn sử dụng', level=1)
    usage_steps = [
        ('Chọn thư mục dữ liệu',
         'click nút "Duyệt..." trên thanh công cụ và chọn thư mục chứa '
         'các file CSV (thư mục data/).'),
        ('Tải dữ liệu',
         'Click nút "Tải dữ liệu" để đọc các file CSV vào bộ nhớ. '
         'Có thể tùy chỉnh số dòng cần đọc bằng trường "Số dòng".'),
        ('Xem thông tin file',
         'Tab "Files" hiển thị danh sách các file đã tải cùng kích thước. '
         'Click vào file để xem dữ liệu trong tab "Dữ liệu".'),
        ('Xem thống kê',
         'Tab "Thống kê" hiển thị thông tin chi tiết về mỗi tập dữ liệu, '
         'bao gồm số dòng, số cột, kiểu dữ liệu, giá trị thiếu.'),
        ('Vẽ biểu đồ',
         'Sử dụng các nút trên thanh biểu đồ hoặc menu "Biểu đồ" để '
         'tạo các loại biểu đồ phân tích.'),
        ('Lưu biểu đồ',
         'Click "Lưu biểu đồ" để xuất biểu đồ hiện tại ra file '
         'PNG, PDF, SVG hoặc JPEG.'),
    ]
    for step_num, (title_text, desc_text) in enumerate(usage_steps, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'Bước {step_num}: {title_text}')
        run.bold = True
        doc.add_paragraph(desc_text)

    # ======= 7. CÁC CHỨC NĂNG CHÍNH =======
    doc.add_heading('7. Các chức năng chính', level=1)

    doc.add_heading('7.1. Tải và quản lý dữ liệu', level=2)
    doc.add_paragraph(
        'Module data_loader.py cung cấp các hàm để đọc từng file CSV riêng '
        'biệt hoặc tải toàn bộ dataset. Hỗ trợ giới hạn số dòng để '
        'preview nhanh với dữ liệu lớn.'
    )

    doc.add_heading('7.2. Phân tích thống kê', level=2)
    doc.add_paragraph(
        'Module analysis.py thực hiện phân tích thống kê mô tả cho các tập dữ liệu:'
    )
    analysis_features = [
        'Thống kê cơ bản: số dòng, cột, kiểu dữ liệu, giá trị thiếu',
        'Phân tích events: phân bổ loại sự kiện, số visitor, số item duy nhất',
        'Phân tích category tree: số danh mục, danh mục gốc vs con',
        'Phân tích item properties: số thuộc tính duy nhất, top thuộc tính',
    ]
    for feat in analysis_features:
        doc.add_paragraph(feat, style='List Bullet')

    doc.add_heading('7.3. Trực quan hóa dữ liệu', level=2)
    doc.add_paragraph(
        'Module visualization.py cung cấp các hàm vẽ biểu đồ:'
    )
    viz_features = [
        'Biểu đồ phân bổ cột (histogram/bar chart) cho mỗi cột dữ liệu',
        'Ma trận tương quan (correlation matrix heatmap) giữa các cột số',
        'Scatter và density plots cho các biến số',
        'Biểu đồ phân bổ sự kiện (bar + pie chart)',
        'Biểu đồ top items được tương tác nhiều nhất',
    ]
    for feat in viz_features:
        doc.add_paragraph(feat, style='List Bullet')

    doc.add_heading('7.4. Giao diện đồ họa (GUI)', level=2)
    doc.add_paragraph(
        'Module gui.py xây dựng giao diện người dùng với Tkinter, bao gồm:'
    )
    gui_features = [
        'Menu bar với các chức năng File, Phân tích, Biểu đồ, Trợ giúp',
        'Toolbar để chọn thư mục dữ liệu và tùy chỉnh số dòng đọc',
        'Panel trái chứa 3 tab: Files, Thống kê, Dữ liệu',
        'Panel phải để hiển thị biểu đồ với toolbar matplotlib tích hợp',
        'Thanh trạng thái (status bar) hiển thị thông tin hoạt động',
        'Chức năng lưu biểu đồ ra file (PNG, PDF, SVG, JPEG)',
    ]
    for feat in gui_features:
        doc.add_paragraph(feat, style='List Bullet')

    # ======= 8. MÔ TẢ KỸ THUẬT =======
    doc.add_heading('8. Mô tả kỹ thuật', level=1)

    doc.add_heading('8.1. Kiến trúc ứng dụng', level=2)
    doc.add_paragraph(
        'Ứng dụng được thiết kế theo kiến trúc module, chia thành các '
        'tầng rõ ràng:'
    )
    arch = [
        'Tầng dữ liệu (data_loader.py): Chịu trách nhiệm đọc và tiền xử lý '
        'các file CSV.',
        'Tầng logic (analysis.py): Thực hiện các phép phân tích thống kê '
        'trên dữ liệu.',
        'Tầng trình bày (visualization.py): Tạo các biểu đồ trực quan hóa '
        'bằng matplotlib.',
        'Tầng giao diện (gui.py): Quản lý giao diện người dùng và tương tác '
        'bằng tkinter.',
    ]
    for item in arch:
        doc.add_paragraph(item, style='List Number')

    doc.add_heading('8.2. Luồng xử lý', level=2)
    doc.add_paragraph(
        '1. Người dùng chọn thư mục chứa dữ liệu CSV thông qua GUI.\n'
        '2. Ứng dụng đọc các file CSV bằng pandas DataFrame.\n'
        '3. Dữ liệu được phân tích thống kê và hiển thị trên tab Thống kê.\n'
        '4. Người dùng có thể xem dữ liệu thô trong tab Dữ liệu.\n'
        '5. Người dùng chọn loại biểu đồ cần vẽ.\n'
        '6. Biểu đồ được tạo bằng matplotlib và hiển thị trên panel phải.\n'
        '7. Người dùng có thể lưu biểu đồ ra file.'
    )

    # ======= 9. KẾT LUẬN =======
    doc.add_heading('9. Kết luận', level=1)
    doc.add_paragraph(
        'Dự án Retailrocket Recommender System đã thành công trong việc '
        'chuyển đổi một Jupyter Notebook thành một ứng dụng Python có '
        'cấu trúc rõ ràng với giao diện đồ họa. Ứng dụng cho phép '
        'người dùng khám phá và phân tích dữ liệu thương mại điện tử '
        'một cách trực quan và hiệu quả.'
    )
    doc.add_paragraph(
        'Các hướng phát triển tiếp theo có thể bao gồm:'
    )
    future = [
        'Xây dựng mô hình gợi ý sản phẩm (collaborative filtering, '
        'content-based filtering)',
        'Phân tích chuỗi thời gian cho dữ liệu sự kiện',
        'Phân khúc khách hàng (customer segmentation) bằng clustering',
        'Tích hợp xuất báo cáo tự động dạng PDF/DOCX',
        'Triển khai ứng dụng web với Flask/Django',
    ]
    for item in future:
        doc.add_paragraph(item, style='List Bullet')

    # ======= LƯU FILE =======
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                                'output', 'Retailrocket_Recommender_Documentation.docx')
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Đã tạo tài liệu thành công: {output_path}")
    return output_path


if __name__ == "__main__":
    create_documentation()
